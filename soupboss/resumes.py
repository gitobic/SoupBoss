"""
Resume management module for SoupBoss.

Handles PDF and text file parsing, content extraction, and storage.
"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import hashlib

from .db import SoupBossDB
from rich.console import Console

console = Console()


class ResumeProcessor:
    """Handles resume file processing and content extraction."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    
    def __init__(self, db: SoupBossDB):
        self.db = db
    
    def process_resume_file(self, file_path: str, name: Optional[str] = None) -> Optional[int]:
        """
        Process a resume file and store it in the database.
        
        Args:
            file_path: Path to the resume file
            name: Optional name for the resume (defaults to filename)
            
        Returns:
            Resume ID if successful, None otherwise
        """
        path = Path(file_path)
        
        # Validate file
        validation_error = self._validate_file(path)
        if validation_error:
            console.print(f"[red]Validation error: {validation_error}[/red]")
            return None
        
        # Use filename as default name if not provided
        if not name:
            name = path.stem
        
        try:
            # Extract text content
            content_text = self._extract_text_content(path)
            if not content_text.strip():
                console.print(f"[red]No text content could be extracted from {file_path}[/red]")
                return None
            
            # Get file metadata
            file_type = path.suffix.lower()[1:]  # Remove the dot
            file_size = path.stat().st_size
            
            # Store in database
            resume_id = self.db.add_resume(
                name=name,
                file_path=str(path.absolute()),
                content_text=content_text,
                file_type=file_type,
                file_size=file_size
            )
            
            console.print(f"[green]Successfully processed resume: {name} (ID: {resume_id})[/green]")
            console.print(f"[dim]Extracted {len(content_text)} characters of text content[/dim]")
            
            return resume_id
            
        except Exception as e:
            console.print(f"[red]Error processing resume {file_path}: {e}[/red]")
            return None
    
    def _validate_file(self, path: Path) -> Optional[str]:
        """Validate resume file. Returns error message if invalid, None if valid."""
        if not path.exists():
            return f"File does not exist: {path}"
        
        if not path.is_file():
            return f"Path is not a file: {path}"
        
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            return f"Unsupported file type: {path.suffix}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            return f"File too large: {file_size / (1024*1024):.1f}MB (max: {self.MAX_FILE_SIZE / (1024*1024):.1f}MB)"
        
        if file_size == 0:
            return "File is empty"
        
        return None
    
    def _extract_text_content(self, path: Path) -> str:
        """Extract text content from resume file based on file type."""
        file_extension = path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(path)
        elif file_extension == '.docx':
            return self._extract_docx_text(path)
        elif file_extension in ['.txt', '.md']:
            return self._extract_text_file(path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from PDF file using PyPDF2."""
        try:
            import PyPDF2
            
            text_content = []
            
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        console.print(f"[yellow]Warning: Could not extract text from page {page_num + 1}: {e}[/yellow]")
                        continue
            
            combined_text = '\n\n'.join(text_content)
            
            # Clean up the text
            combined_text = self._clean_extracted_text(combined_text)
            
            return combined_text
            
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: uv add PyPDF2")
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {e}")
    
    def _extract_docx_text(self, path: Path) -> str:
        """Extract text from DOCX file using python-docx."""
        try:
            from docx import Document
            
            doc = Document(path)
            text_content = []
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            combined_text = '\n'.join(text_content)
            
            # Clean up the text
            combined_text = self._clean_extracted_text(combined_text)
            
            return combined_text
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: uv add python-docx")
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {e}")
    
    def _extract_text_file(self, path: Path) -> str:
        """Extract text from plain text or markdown file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return self._clean_extracted_text(content)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading text file: {e}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text with enhanced PDF artifact handling."""
        import re
        
        # Step 1: Fix common PDF extraction artifacts
        # Fix broken phone numbers like "407 -608-2358" -> "407-608-2358"
        text = re.sub(r'(\d+)\s+(-\d+)', r'\1\2', text)
        
        # Fix broken email domains like "git obic" -> "gitobic"
        text = re.sub(r'(github\.com/\w+)\s+(\w+)', r'\1\2', text)
        text = re.sub(r'(\w+@\w+)\s+(\.com)', r'\1\2', text)
        
        # Fix broken hyphenated words like "on -time" -> "on-time"
        text = re.sub(r'(\w+)\s*-\s*(\w+)', r'\1-\2', text)
        
        # Fix broken compound words with spaces like "enterprise -level" -> "enterprise-level"
        text = re.sub(r'(\w+)\s*-\s*(\w+)', r'\1-\2', text)
        
        # Fix spaces before common punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        
        # Step 2: Normalize whitespace and line breaks
        lines = []
        for line in text.split('\n'):
            # Normalize internal whitespace within lines
            cleaned_line = ' '.join(line.split())
            if cleaned_line:  # Skip empty lines
                lines.append(cleaned_line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(lines)
        
        # Step 3: Final cleanup
        # Remove excessive newlines (max 2 consecutive)
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        # Remove extra spaces around common separators
        cleaned_text = re.sub(r'\s*·\s*', ' · ', cleaned_text)  # Normalize bullet separators
        cleaned_text = re.sub(r'\s*\|\s*', ' | ', cleaned_text)  # Normalize pipe separators
        
        # Fix common broken URLs and email patterns
        cleaned_text = re.sub(r'(\w+)\s+(\.\w+)', r'\1\2', cleaned_text)  # Fix domains
        cleaned_text = re.sub(r'(https?://\w+)\s+(\.\w+)', r'\1\2', cleaned_text)  # Fix URLs
        
        return cleaned_text.strip()
    
    def get_resume_preview(self, resume_id: int, max_length: int = 500) -> Optional[str]:
        """Get a preview of resume content."""
        resume = self.db.get_resume(resume_id)
        if not resume:
            return None
        
        content = resume['content_text']
        if len(content) <= max_length:
            return content
        
        # Find a good break point near the limit
        preview = content[:max_length]
        last_sentence = preview.rfind('.')
        last_newline = preview.rfind('\n')
        
        break_point = max(last_sentence, last_newline)
        if break_point > max_length * 0.7:  # If we have a good break point
            preview = content[:break_point + 1]
        
        return preview + "..." if len(content) > len(preview) else preview
    
    def update_resume_name(self, resume_id: int, new_name: str) -> bool:
        """Update resume name."""
        try:
            cursor = self.db.conn.cursor()
            cursor.execute(
                "UPDATE resumes SET name = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (new_name, resume_id)
            )
            self.db.conn.commit()
            return cursor.rowcount > 0
        except Exception as e:
            console.print(f"[red]Error updating resume name: {e}[/red]")
            return False
    
    def get_file_info(self, file_path: str) -> Dict:
        """Get basic file information without processing."""
        path = Path(file_path)
        
        if not path.exists():
            return {"error": "File not found"}
        
        stat = path.stat()
        return {
            "name": path.name,
            "size": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "extension": path.suffix.lower(),
            "supported": path.suffix.lower() in self.SUPPORTED_EXTENSIONS,
            "too_large": stat.st_size > self.MAX_FILE_SIZE
        }


class ResumeManager:
    """High-level resume management interface."""
    
    def __init__(self, db_path: str = "data/soupboss.db"):
        from .db import get_db
        self.db = get_db(db_path)
        self.processor = ResumeProcessor(self.db)
    
    def add_resume(self, file_path: str, name: Optional[str] = None) -> Optional[int]:
        """Add a resume file to the system."""
        return self.processor.process_resume_file(file_path, name)
    
    def list_resumes(self) -> List[Dict]:
        """Get all resumes with metadata."""
        return self.db.get_resumes()
    
    def get_resume(self, resume_id: int) -> Optional[Dict]:
        """Get specific resume."""
        return self.db.get_resume(resume_id)
    
    def get_resume_preview(self, resume_id: int, max_length: int = 500) -> Optional[str]:
        """Get resume content preview."""
        return self.processor.get_resume_preview(resume_id, max_length)
    
    def remove_resume(self, resume_id: int) -> bool:
        """Remove resume from system."""
        return self.db.delete_resume(resume_id)
    
    def update_resume_name(self, resume_id: int, new_name: str) -> bool:
        """Update resume name."""
        return self.processor.update_resume_name(resume_id, new_name)
    
    def get_file_info(self, file_path: str) -> Dict:
        """Get file information."""
        return self.processor.get_file_info(file_path)
    
    def get_stats(self) -> Dict:
        """Get resume statistics."""
        resumes = self.list_resumes()
        
        if not resumes:
            return {
                "total": 0,
                "by_type": {},
                "total_size_mb": 0
            }
        
        by_type = {}
        total_size = 0
        
        for resume in resumes:
            file_type = resume['file_type']
            by_type[file_type] = by_type.get(file_type, 0) + 1
            total_size += resume['file_size']
        
        return {
            "total": len(resumes),
            "by_type": by_type,
            "total_size_mb": round(total_size / (1024 * 1024), 2)
        }


def get_resume_manager(db_path: str = "data/soupboss.db") -> ResumeManager:
    """Get resume manager instance."""
    return ResumeManager(db_path)